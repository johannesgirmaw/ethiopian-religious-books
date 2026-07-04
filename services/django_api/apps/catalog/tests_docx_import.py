"""Tests for the Word (.docx) book importer and its detection strategies."""

from io import BytesIO

from django.contrib.auth import get_user_model
from django.core.files.uploadedfile import SimpleUploadedFile
from django.test import TestCase
from rest_framework.test import APIClient

from apps.catalog.docx_import import (
    DocxImportError,
    build_chapters_draft_from_docx,
    summarize_docx_structure,
)
from apps.catalog.models import Book
from apps.catalog.pagination import PAGE_BYTE_BUDGET


def _docx(build) -> bytes:
    """Build an in-memory .docx; ``build(doc)`` adds content, returns the bytes."""
    from docx import Document

    doc = Document()
    build(doc)
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _titles(draft):
    return [c["title"] for c in draft]


class DocxStrategyTests(TestCase):
    def test_heading_styles(self):
        def build(doc):
            doc.add_heading("Chapter One", level=1)
            doc.add_paragraph("Body one.")
            doc.add_heading("Chapter Two", level=1)
            doc.add_paragraph("Body two.")

        draft, stats = build_chapters_draft_from_docx(_docx(build), mode="heading")
        self.assertEqual(_titles(draft), ["Chapter One", "Chapter Two"])
        self.assertEqual(stats["strategy_used"], "heading")

    def test_text_patterns_english_and_amharic(self):
        def build(doc):
            doc.add_paragraph("Chapter 1")
            doc.add_paragraph("English body.")
            doc.add_paragraph("ምዕራፍ 2")
            doc.add_paragraph("Amharic body.")

        draft, _ = build_chapters_draft_from_docx(_docx(build), mode="patterns")
        self.assertEqual(_titles(draft), ["Chapter 1", "ምዕራፍ 2"])

    def test_custom_pattern(self):
        def build(doc):
            doc.add_paragraph("§ Alpha")
            doc.add_paragraph("Body A.")
            doc.add_paragraph("§ Beta")
            doc.add_paragraph("Body B.")

        draft, _ = build_chapters_draft_from_docx(
            _docx(build), mode="patterns", pattern=r"^§\s+"
        )
        self.assertEqual(_titles(draft), ["§ Alpha", "§ Beta"])

    def test_format_bold_and_centered(self):
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        def build(doc):
            p1 = doc.add_paragraph()
            r1 = p1.add_run("The Beginning")
            r1.bold = True
            doc.add_paragraph("Some ordinary body text that is clearly long enough.")
            p2 = doc.add_paragraph("A Centered Title")
            p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph("More ordinary body text here for the second chapter.")

        draft, _ = build_chapters_draft_from_docx(_docx(build), mode="format")
        self.assertEqual(_titles(draft), ["The Beginning", "A Centered Title"])

    def test_format_does_not_oversplit_amharic(self):
        # Amharic has no letter case, so plain short lines must NOT be treated as
        # ALL-CAPS titles by the format strategy.
        def build(doc):
            doc.add_paragraph("ጰውሎስ")  # short, not bold/centered
            doc.add_paragraph("የመጽሐፉ ሙሉ ጽሑፍ እዚህ ጋር ይቀጥላል ረጅም ዓረፍተ ነገር ነው።")
            doc.add_paragraph("ጴጥሮስ")
            doc.add_paragraph("ሌላ ረጅም የመጽሐፍ ጽሑፍ እዚህ ጋር ይቀጥላል ዓረፍተ ነገር ነው።")

        draft, _ = build_chapters_draft_from_docx(_docx(build), mode="format")
        # No bold/centered/caps → no chapters detected (single fallback chapter).
        self.assertLessEqual(len(draft), 1)

    def test_page_breaks(self):
        def build(doc):
            doc.add_paragraph("First page intro.")
            p = doc.add_paragraph("Second Chapter")
            p.paragraph_format.page_break_before = True
            doc.add_paragraph("Second chapter body.")

        draft, _ = build_chapters_draft_from_docx(_docx(build), mode="pagebreak")
        self.assertGreaterEqual(len(draft), 2)

    def test_marker_lines(self):
        def build(doc):
            doc.add_paragraph("### Genesis")
            doc.add_paragraph("In the beginning.")
            doc.add_paragraph("<<<CHAPTER: Exodus>>>")
            doc.add_paragraph("These are the names.")

        draft, _ = build_chapters_draft_from_docx(_docx(build), mode="marker")
        self.assertEqual(_titles(draft), ["Genesis", "Exodus"])

    def test_size_fallback_single_chapter(self):
        def build(doc):
            for i in range(5):
                doc.add_paragraph(f"Plain paragraph number {i} with no structure at all.")

        draft, stats = build_chapters_draft_from_docx(_docx(build), mode="size")
        self.assertEqual(len(draft), 1)
        self.assertEqual(stats["strategy_used"], "size")

    def test_auto_prefers_headings_then_falls_back(self):
        def build_headings(doc):
            doc.add_heading("H1", level=1)
            doc.add_paragraph("a")
            doc.add_heading("H2", level=1)
            doc.add_paragraph("b")

        _, stats = build_chapters_draft_from_docx(_docx(build_headings), mode="auto")
        self.assertEqual(stats["strategy_used"], "heading")

        def build_flat(doc):
            doc.add_paragraph("no structure here at all, just one long flat paragraph.")

        _, stats2 = build_chapters_draft_from_docx(_docx(build_flat), mode="auto")
        self.assertEqual(stats2["strategy_used"], "size")

    def test_long_paragraph_split_under_byte_budget(self):
        long_para = "amharic word ደግ " * 400

        def build(doc):
            doc.add_heading("Big", level=1)
            doc.add_paragraph(long_para)

        draft, _ = build_chapters_draft_from_docx(_docx(build), mode="heading")
        pages = draft[0]["pages"]
        self.assertGreater(len(pages), 1)
        for page in pages:
            self.assertLessEqual(len(page["body"].encode("utf-8")), PAGE_BYTE_BUDGET)

    def test_summary_reports_all_modes(self):
        def build(doc):
            doc.add_paragraph("Chapter 1")
            doc.add_paragraph("Body.")
            doc.add_paragraph("Chapter 2")
            doc.add_paragraph("Body.")

        summary = summarize_docx_structure(_docx(build))
        by_mode = {m["mode"]: m for m in summary["modes"]}
        self.assertIn("patterns", by_mode)
        self.assertEqual(by_mode["patterns"]["chapters"], 2)
        self.assertEqual(summary["recommended"], "patterns")
        self.assertEqual(by_mode["patterns"]["titles"], ["Chapter 1", "Chapter 2"])

    def test_invalid_file_raises(self):
        with self.assertRaises(DocxImportError):
            build_chapters_draft_from_docx(b"not a real docx")

    def test_bad_regex_raises_value_error(self):
        with self.assertRaises(ValueError):
            build_chapters_draft_from_docx(_docx(lambda d: d.add_paragraph("x")),
                                           mode="patterns", pattern="(")


class DocxImportApiTests(TestCase):
    def setUp(self):
        self.author = get_user_model().objects.create_user(
            email="author@example.com", password="pw123456", role="author"
        )
        self.client = APIClient()
        self.client.force_authenticate(self.author)

    def _upload(self, name="book.docx"):
        def build(doc):
            doc.add_heading("Chapter One", level=1)
            doc.add_paragraph("Body one.")
            doc.add_heading("Chapter Two", level=1)
            doc.add_paragraph("Body two.")

        return SimpleUploadedFile(
            name,
            _docx(build),
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

    def test_import_creates_draft_book(self):
        res = self.client.post(
            "/v1/admin/books/import-docx",
            {"file": self._upload(), "primary_language": "am", "mode": "auto"},
            format="multipart",
        )
        self.assertEqual(res.status_code, 201, res.data)
        self.assertEqual(res.data["catalog_visibility"], "hidden")
        self.assertEqual(len(res.data["chapters_draft"]), 2)
        self.assertEqual(res.data["import_stats"]["strategy_used"], "heading")
        book = Book.objects.get(id=res.data["id"])
        self.assertEqual(book.created_by_id, self.author.id)
        self.assertEqual(book.title, "Chapter One")

    def test_preview_returns_per_mode_counts(self):
        res = self.client.post(
            "/v1/admin/books/import-docx/preview",
            {"file": self._upload()},
            format="multipart",
        )
        self.assertEqual(res.status_code, 200, res.data)
        self.assertIn("modes", res.data)
        self.assertIn("recommended", res.data)
        by_mode = {m["mode"]: m for m in res.data["modes"]}
        self.assertEqual(by_mode["heading"]["chapters"], 2)

    def test_legacy_doc_rejected_with_guidance(self):
        bad = SimpleUploadedFile("old.doc", b"\xd0\xcf\x11\xe0stuff", content_type="application/msword")
        res = self.client.post(
            "/v1/admin/books/import-docx", {"file": bad}, format="multipart"
        )
        self.assertEqual(res.status_code, 400)
        self.assertEqual(res.data["error"]["code"], "LEGACY_DOC_FORMAT")

    def test_bad_custom_pattern_returns_400(self):
        res = self.client.post(
            "/v1/admin/books/import-docx",
            {"file": self._upload(), "mode": "patterns", "pattern": "("},
            format="multipart",
        )
        self.assertEqual(res.status_code, 400)
